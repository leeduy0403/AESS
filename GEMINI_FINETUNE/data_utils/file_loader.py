import requests, mimetypes
from io import BytesIO
from docx import Document
import pdfplumber

def load_file_content(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        content_type = response.headers.get("Content-Type", "").lower()
        ext = mimetypes.guess_extension(content_type.split(";")[0]) or ""
        if ".pdf" in ext or "pdf" in content_type:
            return extract_pdf_text(BytesIO(response.content))
        elif ".docx" in ext or "wordprocessingml" in content_type:
            doc = Document(BytesIO(response.content))
            return "\n".join([p.text for p in doc.paragraphs])
        elif "text/plain" in content_type or ext == ".txt":
            return response.text
        else:
            return f"Unsupported file type from URL: {content_type}"
    except Exception as e:
        return f"Error loading file from URL: {e}"

def extract_pdf_text(pdf_path):
    output = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            height = page.height
            cropped = page.within_bbox((0, 50, page.width, height - 50))
            output.append(cropped.extract_text())
    return "\n\n".join(output)

def extract_docx_text(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error extracting DOCX: {e}"