import re 
import os
import time
import mimetypes
import requests
import json
import google.generativeai as genai
import pandas as pd
# from load_creds import load_creds
from io import BytesIO
from pdfminer.high_level import extract_text
import pdfplumber
from docx import Document
from vertexai.generative_models import GenerativeModel
from sklearn.metrics import cohen_kappa_score
from dotenv import load_dotenv

load_dotenv()

model = GenerativeModel(
    model_name="projects/558798320545/locations/us-central1/endpoints/7445388067461922816", 
    # model_name="projects/558798320545/locations/us-central1/endpoints/816652365925974016",
    generation_config={
        "temperature": 0.3,
        "top_p": 0.95,
        "max_output_tokens": 8192,
    }
)

feedback_model = genai.GenerativeModel("gemini-2.0-flash")

def load_file_content(url):
    """Download and extract text content from a file at a given URL (PDF, DOCX, or TXT)."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        content_type = response.headers.get("Content-Type", "").lower()
        ext = mimetypes.guess_extension(content_type.split(";")[0]) or ""
        if ".pdf" in ext or "pdf" in content_type:
            return extract_pdf_text(BytesIO(response.content))
        elif ".docx" in ext or "wordprocessingml" in content_type:
            doc = Document(BytesIO(response.content))
            return "\n".join([p.text for p in doc.paragraphs])
        elif "text/plain" in content_type or ext == ".txt":
            return response.text
        else:
            return f"Unsupported file type from URL: {content_type}"
    except Exception as e:
        return f"Error loading file from URL: {e}"

def extract_pdf_text(pdf_path):
    output = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            height = page.height
            # Crop to remove top and bottom 50 pts
            cropped = page.within_bbox((0, 50, page.width, height - 50))
            output.append(cropped.extract_text())
    return "\n\n".join(output)
    
def extract_docx_text(file_path):
    """Extract text from a DOCX file."""
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error extracting DOCX: {e}"

def normalize_number(text):
    """Converts a string with ',' or '.' to a float, and handles trailing dot like '2.'."""
    return float(text.strip().replace(',', '.').rstrip('.'))

def extract_score_ranges_and_components(description_text):
    """Extracts score ranges, components, and coefficients from the DESCRIPTION text."""

    float_pattern = r"\d+(?:[.,]\d+)?\.?"

    # Match total score range
    total_score_match = re.search(
        rf"Total\s+Scores?\s+range:\s*({float_pattern})\s*-\s*({float_pattern})[\.]", 
        description_text, 
        re.IGNORECASE
    )
    if not total_score_match:
        total_score_match = re.search(
            rf"Scores?\s+range:\s*({float_pattern})\s*-\s*({float_pattern})[\.]", 
            description_text, 
            re.IGNORECASE
        )

    # Match component score range
    comp_score_match = re.search(
        rf"Components?\s+Scores?\s+range:\s*({float_pattern})\s*-\s*({float_pattern})[\.]", 
        description_text, 
        re.IGNORECASE
    )

    # Match components
    components_match = re.search(r"Components?:\s*([^\.]+)", description_text, re.IGNORECASE)

    # Match coefficients (now more robust)
    coefficients_match = re.search(
        r"Coefficients?:\s*([0-9.,;\s]+)", description_text, re.IGNORECASE
    )
    if not coefficients_match:
        coefficients_match = re.search(
            r"Coefficient:\s*([0-9.,;\s]+)", description_text, re.IGNORECASE
        )

    # Extract total score range
    if total_score_match:
        min_total, max_total = map(normalize_number, total_score_match.groups())
    else:
        min_total, max_total = 0.0, 0.0

    # Extract component score range
    if comp_score_match:
        min_comp, max_comp = map(normalize_number, comp_score_match.groups())
    else:
        min_comp, max_comp = 0.0, 0.0

    # Extract components
    components = []
    if components_match:
        components_text = components_match.group(1)
        components = [comp.strip() for comp in components_text.split(',') if comp.strip()]

    # Extract coefficients (split by `;`)
    coefficients = []
    if coefficients_match:
        coefficients_text = coefficients_match.group(1)
        coefficients = [
            normalize_number(coef) for coef in re.split(r';\s*', coefficients_text) if coef.strip()
        ]

    # Mismatch check
    if components and coefficients and len(components) != len(coefficients):
        print(f"⚠️ Warning: {len(components)} components but {len(coefficients)} coefficients found!")

    print(f"Total Score Range: {min_total}-{max_total}, Component Score Range: {min_comp}-{max_comp}, Components: {components}, Coefficients: {coefficients}")
    return (min_total, max_total), (min_comp, max_comp), components, coefficients

def calculate_weighted_kappa(y_true, y_pred):
    """Compute quadratic weighted kappa score."""
    try:
        kappa = cohen_kappa_score(y_true, y_pred, weights='quadratic')
    except Exception as e:
        print(f"⚠️ Error calculating kappa: {e}")
        kappa = 0.0
    return kappa

def clean_feedback(feedback):
    """Cleans feedback by removing inline numbers and fixing abrupt endings."""
    # Remove markdown-style bold (**text** or __text__)
    feedback = re.sub(r"(\*\*|__)(.*?)\1", r"\2", feedback)
    # Remove inline numbers (like 'example1', 'point2', but not section numbers like 1. Introduction)
    feedback = re.sub(r"(?<!\s)\d+", "", feedback)
    return feedback.strip()

def extract_score(response_text, min_total, max_total, min_comp, max_comp, components, coefficients):
    """Extracts overall score, individual component scores and feedback."""

    score = min_total
    component_scores = [min_comp] * len(components)
    lines = response_text.strip().split("\n")
    
    for line in lines:
        line = line.strip()
        # Handle total score extraction
        if line.lower().startswith("score:"):
            match = re.search(r"\d+(?:[.,]\d+)?", line)
            if match:
                score_val = float(match.group(0).replace(',', '.'))
                score = min(max_total, max(min_total, score_val))

        # Handle component scores
        for comp in components:
            pattern = re.compile(rf"{re.escape(comp)}\s*:\s*(\d+(?:[.,]\d+)?)", re.IGNORECASE)
            match = pattern.search(response_text)
            if match:
                score_val = float(match.group(1).replace(',', '.'))
                component_scores[components.index(comp)] = min(max_comp, max(min_comp, score_val))

    # If component scores are extracted, recalculate total
    if components:
    #     if coefficients:
    #         score = sum(coeff * comp_score for coeff, comp_score in zip(coefficients, component_scores))
    #     else:
    #         score = sum(component_scores)
        if coefficients:
            score = 0
            for i, (coeff, comp_score) in enumerate(zip(coefficients, component_scores)):
                partial = coeff * comp_score
                print(f"Component {i+1}: Coefficient={coeff}, Score={comp_score} → Partial={partial}")
                score += partial
        else:
            score = sum(component_scores)
    print(f"Final Score: {score}")
    return score, component_scores

def evaluate_essays(description_path, content_path, output_path):
    """Evaluates essays based on the extracted description and generates scores & feedback."""
    if description_path.lower().startswith("http"):
        description_content = load_file_content(description_path)
    else:
        if description_path.lower().endswith(".pdf"):
            description_content = extract_pdf_text(description_path)
        elif description_path.lower().endswith(".docx"):
            description_content = extract_docx_text(description_path)
        else:
            with open(description_path, encoding="utf-8") as f:
                description_content = f.read()
    (min_total, max_total), (min_comp, max_comp), components, coefficients = extract_score_ranges_and_components(description_content)

    df = pd.read_excel(content_path)
    df = df.head(1).reset_index(drop=True)
    results = []
    correct_count = 0
    total_count = 0

    # Rate limiting variables (30 requests per minute)
    request_count = 0
    start_time = time.time()

    for index, row in df.iterrows():
        essay_id = row["essay_id"]
        essay_content = row["essay"]
        true_score = row.get("rater1_domain1", None)
        # print(f"Processing essay_id {essay_id}")

        try:
            # Rate limiting check
            request_count += 1
            elapsed = time.time() - start_time
            if request_count > 100:
                if elapsed < 60:
                    wait_time = 60 - elapsed
                    print(f"⏳ Rate limit reached. Waiting {wait_time:.2f} seconds...")
                    time.sleep(wait_time)
                request_count = 1
                start_time = time.time()

            # Build the scoring prompt
            prompt = (
                f"Evaluate the given CONTENT based on the DESCRIPTION. For each essay, provide the following:\n"
                f"1. An overall score.\n"
                f"2. A score for each component (if DESCRIPTION has).\n"
                f"DESCRIPTION: {description_content}\n"
                f"CONTENT: {essay_content}\n\n"
                f"Follow the exact format below in your response:\n\n"
                f"Score: (a number from {min_total}-{max_total})\n"
            )
            for comp in components:
                prompt += f"{comp}: (a number from {min_comp}-{max_comp})\n"
            prompt += f"Example:\nScore: {min_total}\n"
            for comp in components:
                prompt += f"{comp}: {min_comp}\n"

            # Feedback prompt
            fb_prompt = (
                f"DESCRIPTION: {description_content}\n"
                f"CONTENT: {essay_content}\n\n"
                f"Your task: Provide **only constructive feedback** on the CONTENT, based on the DESCRIPTION. Please give feedback as many details as possible\n"
                f"Do NOT give any score. Be specific and concise. Mention strengths and areas to improve.\n\n"
                f"If DESCRIPTION has components, provide feedback for each component.\n"
                f"Follow the exact format below in your response:\n\n"
                f"Overall Feedback: (Provide holistic feedback about the essay.)\n"
            )

            for comp in components:
                fb_prompt += f"{comp} Feedback: (Provide specific feedback for the '{comp}' aspect.)\n"

            fb_prompt += (
                f"Example:\n"
                f"Overall Feedback: The essay presents a clear argument with appropriate structure.\n"
            )

            for comp in components:
                fb_prompt += f"{comp} Feedback: Needs more examples.\n"

            # Retry logic for 429 errors (retry up to 5 times)
            max_retries = 5
            for attempt in range(max_retries):
                try:
                    score_rp = model.generate_content(prompt)
                    score_text = score_rp.text.strip()
                    # if not score_text:
                    #     print(f"⚠️ Empty response for essay_id {essay_id}. Retrying... (Attempt {attempt+1}/{max_retries})")
                    #     time.sleep(10)  # Add a short delay before retrying
                    #     continue  # Skip the rest of the loop and retry
                    
                    print(f"Response for essay_id {essay_id}:\n{score_text}")
                    break  # Break out if successful
                except Exception as e:
                    err_msg = str(e)
                    if "429 Resource exhausted" in err_msg:
                        print(f"⚠️ Essay {essay_id}: {err_msg}")
                        print(f"Waiting 60 seconds before retrying (Attempt {attempt+1}/{max_retries})...")
                        time.sleep(60)
                    else:
                        raise  # For any other error, do not retry
            
            # Process the generated score text
            score, scores = extract_score(
                score_text, min_total, max_total, min_comp, max_comp, components, coefficients
            )

            feedback_rp = feedback_model.generate_content(fb_prompt)
            feedback_text = feedback_rp.text.strip()

            feedback = clean_feedback(feedback_text)
            print(f"Feedback for essay_id {essay_id}:\n{feedback}")

            # if true_score is not None and int(true_score) == score:
            #     correct_count += 1
            # total_count += 1

            results.append({
                "essay_id": essay_id,
                "ovr": score,
                "scores": scores,
                "components": components,
                "coefficients": coefficients,
                "feedback": 'Feedback not generated due to error.'
            })
        except Exception as e:
            print(f"Error processing essay_id {essay_id}: {e}")
            results.append({
                "essay_id": essay_id,
                "ovr": min_total,
                "scores": [min_comp] * len(components),
                "components": components,
                "coefficients": coefficients,
                "feedback": "Error processing submission."
            })

    # accuracy = correct_count / total_count if total_count else 0
    # print(f"✅ Accuracy compared to rater1_domain1: {accuracy:.2%} ({correct_count}/{total_count})")

    # true_scores = df["rater1_domain1"].astype(int).tolist()
    # llm_scores = [r["ovr"] for r in results]
    # kappa_llm = calculate_weighted_kappa(true_scores, llm_scores)
    # print(f"🎯 Weighted Kappa (LLM): {kappa_llm:.4f}")

    results_df = pd.DataFrame(results)
    final_df = df.merge(results_df, on="essay_id", how="left")
    final_df.to_excel(output_path, index=False)
    print(f"✅ Evaluation complete. Results saved to {output_path}")


if __name__ == "__main__":
    evaluate_essays(
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET1//essay_set_1_description.docx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET1//essay_set_1_test.xlsx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//OUTPUT//TEST_FLOW//set1.xlsx"
    )
    evaluate_essays(
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET2//essay_set_2_description.docx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET2//essay_set_2_test.xlsx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//OUTPUT//TEST_FLOW//set2.xlsx"
    )
    evaluate_essays(
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET3//essay_set_3_description.docx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET3//essay_set_3_test.xlsx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//OUTPUT//TEST_FLOW//set3.xlsx"
    )
    evaluate_essays(
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET4//essay_set_4_description.docx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET4//essay_set_4_test.xlsx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//OUTPUT//TEST_FLOW//set4.xlsx"
    )
    evaluate_essays(
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET5//essay_set_5_description.docx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET5//essay_set_5_test.xlsx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//OUTPUT//TEST_FLOW//set5.xlsx"
    )
    evaluate_essays(
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET6//essay_set_6_description.docx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET6//essay_set_6_test.xlsx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//OUTPUT//TEST_FLOW//set6.xlsx"
    )
    evaluate_essays(
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET7//essay_set_7_description.docx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET7//essay_set_7_test.xlsx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//OUTPUT//TEST_FLOW//set7.xlsx"
    )
    evaluate_essays(
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET8//essay_set_8_description.docx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//Submission//SET8//essay_set_8_test.xlsx",
        "C://Users//Admin//OneDrive//DACN+DATN//AI MODEL//AES//OUTPUT//TEST_FLOW//set8.xlsx"
    )